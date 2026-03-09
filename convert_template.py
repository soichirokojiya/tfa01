"""
倉元製作所の報告書をテンプレートに変換するスクリプト
- 会社固有の値をプレースホルダーに置換
- 空段落の連続を改ページに修正
"""

from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

SOURCE = "/Users/apple/Library/CloudStorage/GoogleDrive-koujiy@souichirou.org/マイドライブ/TFA/倉元製作所/202602/202602v03_新株予約権算定報告書_株式会社倉元製作所.docx"
OUTPUT = "/Users/apple/tfa01/template.docx"


def replace_in_runs(paragraph, old_text, new_text):
    """段落内のrunを結合してテキスト置換"""
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False
    new_full = full_text.replace(old_text, new_text)
    runs = paragraph.runs
    if not runs:
        return False
    remaining = new_full
    for run in runs:
        run_len = len(run.text)
        if len(remaining) <= 0:
            run.text = ""
        elif len(remaining) <= run_len:
            run.text = remaining
            remaining = ""
        else:
            run.text = remaining[:run_len]
            remaining = remaining[run_len:]
    if remaining:
        runs[-1].text += remaining
    return True


def replace_all(doc, old_text, new_text):
    """ドキュメント全体でテキスト置換"""
    count = 0
    for para in doc.paragraphs:
        if replace_in_runs(para, old_text, new_text):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_runs(para, old_text, new_text):
                        count += 1
    return count


def fix_empty_paragraphs(doc):
    """
    3個以上連続する空段落を削除し、最初の空段落に改ページを設定。
    ただしドキュメント先頭（表紙用）の空段落は保持。
    """
    body = doc.element.body
    paragraphs = doc.paragraphs

    # 空段落の連続区間を特定
    regions = []
    i = 0
    while i < len(paragraphs):
        if paragraphs[i].text.strip() == '':
            start = i
            while i < len(paragraphs) and paragraphs[i].text.strip() == '':
                i += 1
            count = i - start
            if count >= 3:
                # 次の段落のテキスト
                next_text = paragraphs[i].text.strip() if i < len(paragraphs) else ""
                regions.append((start, i - 1, count, next_text))
        else:
            i += 1

    print(f"空段落の連続区間: {len(regions)}箇所")
    for start, end, count, next_text in regions:
        print(f"  段落{start}〜{end} ({count}個) → [{next_text[:40]}]")

    # 後ろから処理（削除してもインデックスがずれないように）
    for start, end, count, next_text in reversed(regions):
        # 先頭4行（表紙レイアウト用）はスキップ
        if start == 0:
            continue

        # 最初の空段落を残し、改ページを設定
        first_p = paragraphs[start]._element

        # 残りの空段落を削除
        for idx in range(end, start, -1):
            elem = paragraphs[idx]._element
            body.remove(elem)

        # 最初の空段落に改ページを追加
        pPr = first_p.find(qn('w:pPr'))
        if pPr is None:
            pPr = first_p.makeelement(qn('w:pPr'), {})
            first_p.insert(0, pPr)

        # 既存のpageBreakBeforeを削除
        for pb in pPr.findall(qn('w:pageBreakBefore')):
            pPr.remove(pb)

        # pageBreakBeforeを追加
        pb = pPr.makeelement(qn('w:pageBreakBefore'), {})
        pPr.append(pb)

    print("改ページ修正完了")


def main():
    doc = Document(SOURCE)

    # ── テキスト置換（順序重要！長い文字列→短い文字列の順）──

    # CAPM式（最も長いので先に）
    c = replace_all(doc, "= 1.702% + 9.3% × 0.777 + 21.83%",
                         "= 1.591% + 9.3%× 0.567 + 21.83%")
    print(f"CAPM式: {c}箇所")

    c = replace_all(doc, "= 30.76%", "= 28.69%")
    print(f"CAPM結果: {c}箇所")

    # 流動性の記述（段落156の長い文）
    c = replace_all(doc, "14,765株（2021年2月10日から2026年2月9日までの日次売買高の中央値である147,650株",
                         "148,313株（2021年3月3日から2026年3月2日までの日次売買高の中央値である1,483,123株")
    print(f"流動性記述: {c}箇所")

    # ボラティリティ期間（表0内）
    c = replace_all(doc, "2021年1月- 2026年1月", "2021年2月- 2026年2月")
    print(f"ボラティリティ期間: {c}箇所")

    # β期間（表0内）
    c = replace_all(doc, "2021年2月10日から2026年2月9日の日次β",
                         "2021年3月3日から2026年3月2日の日次β")
    print(f"β期間: {c}箇所")

    # 権利行使期間（テーブル0とテーブル1に出現）
    # テーブル1の権利行使期間
    t1 = doc.tables[1]
    replace_in_runs(t1.rows[5].cells[1].paragraphs[0],
                    "2026年3月31日- 2031年3月30日",
                    "2026年3月3日- 2026年3月4日")

    # テーブル0の権利行使期間
    t0 = doc.tables[0]
    replace_in_runs(t0.rows[0].cells[1].paragraphs[4],
                    "2026年3月31日- 2031年3月30日",
                    "2026年3月3日- 2026年3月4日")
    print("権利行使期間: 置換完了")

    # テーブル1の決議年月日
    replace_in_runs(t1.rows[0].cells[1].paragraphs[0],
                    "2026年3月31日", "2026年●月●日")
    print("決議年月日: 置換完了")

    # テーブル1の総数・株式数・払込価額
    replace_in_runs(t1.rows[2].cells[1].paragraphs[0], "37,147個", "●個")
    replace_in_runs(t1.rows[3].cells[1].paragraphs[0], "3,714,700株", "●株")
    replace_in_runs(t1.rows[4].cells[1].paragraphs[0], "772,657,600円", "●円")
    print("テーブル1 数値: 置換完了")

    # テーブル0の権利行使価格
    replace_in_runs(t0.rows[0].cells[1].paragraphs[1], "208円", "●円")
    print("権利行使価格: 置換完了")

    # 会社名（先にフルネームを置換）
    c = replace_all(doc, "株式会社倉元製作所", "株式会社ジェリービーンズグループ")
    print(f"会社フルネーム: {c}箇所")
    c = replace_all(doc, "倉元製作所", "ジェリービーンズグループ")
    print(f"会社名: {c}箇所")

    # 証券コード
    c = replace_all(doc, "5216", "3070")
    print(f"証券コード: {c}箇所")

    # 株価
    c = replace_all(doc, "231円", "110円")
    print(f"株価: {c}箇所")

    # ボラティリティ
    c = replace_all(doc, "51.12%", "62.54%")
    print(f"ボラティリティ: {c}箇所")

    # リスクフリーレート（テーブル0内の1.702%）
    c = replace_all(doc, "1.702%", "1.591%")
    print(f"リスクフリーレート: {c}箇所")

    # β値
    c = replace_all(doc, "0.777", "0.567")
    print(f"β値: {c}箇所")

    # 評価基準日
    c = replace_all(doc, "2026年2月9日", "2026年3月2日")
    print(f"評価基準日: {c}箇所")

    # 発行済株式数
    c = replace_all(doc, "47,998,575", "79,440,000")
    print(f"発行済株式数: {c}箇所")

    # 代表者
    c = replace_all(doc, "渡邉\u3000敏行", "宮崎\u3000明")
    print(f"代表者（全角スペース）: {c}箇所")
    c = replace_all(doc, "渡邉　敏行", "宮崎　明")
    print(f"代表者（全角スペース2）: {c}箇所")

    # 住所
    c = replace_all(doc, "宮城県栗原市若柳武鎗字花水前１－１",
                         "東京都台東区上野1-16-5")
    print(f"住所: {c}箇所")

    # 設立年月
    c = replace_all(doc, "1990年8月", "1990年4月")
    print(f"設立年月: {c}箇所")

    # 決算日
    c = replace_all(doc, "12月末", "1月末")
    print(f"決算日: {c}箇所")

    # 国債の償還日表記（テーブル0内）
    c = replace_all(doc, "2031年3月20日償還の国債レート", "2031年3月20日償還の国債レート")
    # 同じなのでスキップ

    # 公正価値（●のまま）- すでに●なのでそのまま

    # 第●回 - すでに●なのでそのまま

    # ── 空段落 → 改ページ修正 ──
    fix_empty_paragraphs(doc)

    # ── 保存 ──
    doc.save(OUTPUT)
    print(f"\n保存完了: {OUTPUT}")


if __name__ == "__main__":
    main()

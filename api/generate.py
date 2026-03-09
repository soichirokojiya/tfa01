"""
Vercel Serverless Function: 新株予約権評価報告書を生成して返す
POST /api/generate
"""

import json
import io
import os
import re
import math
import urllib.request
import urllib.parse
import tempfile
import zipfile
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
from http.server import BaseHTTPRequestHandler
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
import yfinance as yf
import numpy as np
import warnings
warnings.filterwarnings("ignore")

# Vercel: テンプレートファイルのパスを複数候補から解決
_candidates = [
    os.path.join(os.path.dirname(__file__), "..", "template.docx"),
    os.path.join(os.path.dirname(__file__), "template.docx"),
    "/var/task/template.docx",
]
TEMPLATE_PATH = next((p for p in _candidates if os.path.exists(p)), _candidates[0])


def fetch_jsda_bond(eval_dt, exercise_end_dt) -> dict:
    """JSDAサイトから売買参考統計値をDLし、権利行使期間終了に最も近い長期国債を返す"""
    result = {"name": "", "maturity": "", "yield_value": "", "all_bonds": []}
    try:
        import xlrd
        # ファイル名: S + 西暦下2桁 + MMDD
        yy = eval_dt.year % 100
        fname = f"S{yy:02d}{eval_dt.month:02d}{eval_dt.day:02d}"
        url = f"https://market.jsda.or.jp/shijyo/saiken/baibai/baisanchi/files/{eval_dt.year}/{fname}.xls"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = resp.read()
        wb = xlrd.open_workbook(file_contents=data)
        ws = wb.sheet_by_index(0)
        # 長期国債データを収集（「超長期国債」は除外）
        all_bonds = []
        best = None
        for r in range(1, ws.nrows):
            name = str(ws.cell_value(r, 2)).strip()
            if "長期国債" not in name or "超長期国債" in name:
                continue
            maturity_str = str(ws.cell_value(r, 3))
            coupon = ws.cell_value(r, 4)
            # 平均値
            avg_price = ws.cell_value(r, 5)
            avg_change = ws.cell_value(r, 6)
            avg_compound = ws.cell_value(r, 7)
            avg_simple = ws.cell_value(r, 8)
            # 中央値
            med_price = ws.cell_value(r, 9)
            med_change = ws.cell_value(r, 10)
            med_compound = ws.cell_value(r, 11)
            med_simple = ws.cell_value(r, 12)
            if not med_compound:
                continue
            try:
                mat_dt = datetime.strptime(maturity_str, "%Y/%m/%d")
            except (ValueError, TypeError):
                continue
            diff = abs((mat_dt - exercise_end_dt).days)
            bond_entry = {
                "name": name,
                "maturity": mat_dt,
                "coupon": coupon,
                "avg_price": avg_price,
                "avg_change": avg_change,
                "avg_compound": avg_compound,
                "avg_simple": avg_simple,
                "med_price": med_price,
                "med_change": med_change,
                "med_compound": med_compound,
                "med_simple": med_simple,
                "yield_value": float(med_compound),
                "diff_days": diff,
            }
            all_bonds.append(bond_entry)
            if best is None or diff < best[0]:
                best = (diff, name, mat_dt, float(med_compound))
        # 償還日順にソート
        all_bonds.sort(key=lambda x: x["maturity"])
        result["all_bonds"] = all_bonds
        if best:
            result["name"] = best[1]
            result["maturity"] = best[2].strftime("%Y-%m-%d")
            result["yield_value"] = str(best[3])
    except Exception:
        pass
    return result


def fetch_yahoo_quote_data(ticker_code: str) -> dict:
    """Yahoo Finance Japan から発行済株式数・配当情報を取得"""
    result = {"shares_outstanding": 0, "dividend_per_share": 0, "dividend_yield": 0.0}
    try:
        url = f"https://finance.yahoo.co.jp/quote/{ticker_code}.T"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            html = resp.read().decode("utf-8")
        # HTML から直接スクレイピング
        m = re.search(
            r'発行済株式数.*?<span[^>]*class="StyledNumber__value[^"]*"[^>]*>([\d,]+)</span>',
            html, re.DOTALL)
        if m:
            result["shares_outstanding"] = int(m.group(1).replace(",", ""))
        m2 = re.search(
            r'配当利回り.*?<span[^>]*class="StyledNumber__value[^"]*"[^>]*>([\d.]+)</span>',
            html, re.DOTALL)
        if m2:
            result["dividend_yield"] = round(float(m2.group(1)), 2)
        m3 = re.search(
            r'1株配当.*?<span[^>]*class="StyledNumber__value[^"]*"[^>]*>([\d.]+)</span>',
            html, re.DOTALL)
        if m3:
            result["dividend_per_share"] = int(float(m3.group(1)))
    except Exception:
        pass

    # フォールバック: yfinance から発行済株式数を取得
    if result["shares_outstanding"] == 0:
        try:
            ticker = yf.Ticker(f"{ticker_code}.T")
            info = ticker.info
            shares = info.get("sharesOutstanding", 0)
            if shares:
                result["shares_outstanding"] = int(shares)
        except Exception:
            pass

    return result


# ──────────────────────────────────────────────
# データ取得
# ──────────────────────────────────────────────

def fetch_japanese_company_name(ticker_code: str) -> str:
    # 1) Yahoo Finance Japan Webページ
    try:
        url = f"https://finance.yahoo.co.jp/quote/{ticker_code}.T"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            html = resp.read().decode("utf-8")
        m = re.search(r"<title>(.*?)【\d+】", html)
        if m:
            name = m.group(1).strip()
            name = re.sub(r"^\(株\)", "", name)
            name = re.sub(r"\(株\)$", "", name)
            name = re.sub(r"^株式会社", "", name)
            name = re.sub(r"株式会社$", "", name)
            return name.strip()
    except Exception:
        pass

    # 2) フォールバック: yfinance API
    try:
        ticker = yf.Ticker(f"{ticker_code}.T")
        info = ticker.info
        name = info.get("shortName", "") or info.get("longName", "")
        name = re.sub(r"^\(株\)", "", name)
        name = re.sub(r"\(株\)$", "", name)
        name = re.sub(r"^株式会社", "", name)
        name = re.sub(r"株式会社$", "", name)
        return name.strip()
    except Exception:
        pass

    raise ValueError(f"社名を取得できませんでした: {ticker_code}")


def fetch_company_profile(ticker_code: str) -> dict:
    profile = {"representative": "", "address": "", "established": "", "settlement": ""}

    # 1) Yahoo Finance Japan Webページ
    try:
        url = f"https://finance.yahoo.co.jp/quote/{ticker_code}.T/profile"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            html = resp.read().decode("utf-8")

        def extract(label):
            m = re.search(rf'<th[^>]*>{label}</th>\s*<td[^>]*>(.*?)</td>', html, re.DOTALL)
            if m:
                return re.sub(r'<[^>]+>', '', m.group(1)).strip()
            return ""

        m = re.search(r'〒[\d\-]+\s*(.+?)(?=<|")', html)
        profile["representative"] = extract("代表者名")
        profile["address"] = m.group(1).strip() if m else ""
        profile["established"] = extract("設立年月日")
        profile["settlement"] = extract("決算")
    except Exception:
        pass

    return profile


def fetch_stock_data(ticker_code: str, eval_date: str, exercise_end: str = ""):
    ticker = yf.Ticker(f"{ticker_code}.T")
    eval_dt = datetime.strptime(eval_date, "%Y-%m-%d")

    # 権利行使期間終了日から期間を算出
    if exercise_end:
        ex_end_dt = datetime.strptime(exercise_end, "%Y-%m-%d")
    else:
        # デフォルト: 5年
        ex_end_dt = eval_dt + relativedelta(years=5)

    # 基準日から満期までの月数（ボラティリティ用）
    rd = relativedelta(ex_end_dt, eval_dt)
    months_to_maturity = rd.years * 12 + rd.months
    if rd.days > 0:
        months_to_maturity += 1  # 端数月は切り上げ

    # 基準日から満期までの日数（出来高用）
    days_to_maturity = (ex_end_dt - eval_dt).days

    start = (eval_dt - timedelta(days=10)).strftime("%Y-%m-%d")
    end = (eval_dt + timedelta(days=1)).strftime("%Y-%m-%d")
    hist_around = ticker.history(start=start, end=end)
    hist_before = hist_around[hist_around.index.strftime("%Y-%m-%d") <= eval_date]
    if len(hist_before) == 0:
        raise ValueError(f"評価基準日 {eval_date} の株価データが取得できません")
    stock_price = int(hist_before["Close"].iloc[-1])

    # ボラティリティ: 基準日の前月から months_to_maturity ヶ月遡る
    vol_end_month = eval_dt.replace(day=1) - timedelta(days=1)  # 前月末
    vol_start_month = vol_end_month - relativedelta(months=months_to_maturity - 1)
    vol_start = vol_start_month.replace(day=1).strftime("%Y-%m-%d")
    vol_end = (vol_end_month + timedelta(days=1)).strftime("%Y-%m-%d")
    hist_monthly = ticker.history(start=vol_start, end=vol_end, interval="1mo")
    returns = np.log(hist_monthly["Close"] / hist_monthly["Close"].shift(1)).dropna()
    annual_vol = returns.std(ddof=0) * np.sqrt(12)
    vol_start_label = f"{vol_start_month.year}年{vol_start_month.month}月"
    vol_end_label = f"{vol_end_month.year}年{vol_end_month.month}月"

    # 出来高: 基準日から days_to_maturity 日分を遡る
    report_date = eval_dt
    volume_end = eval_dt
    volume_start = eval_dt - timedelta(days=days_to_maturity)
    hist_daily = ticker.history(
        start=volume_start.strftime("%Y-%m-%d"),
        end=(volume_end + timedelta(days=1)).strftime("%Y-%m-%d"),
    )
    median_volume = int(hist_daily["Volume"].median())
    liquidity_shares = math.ceil(median_volume * 0.1)

    # Yahoo Finance Japan から配当・発行済株式数を取得
    yahoo_data = fetch_yahoo_quote_data(ticker_code)
    dividend_per_share = yahoo_data["dividend_per_share"]
    dividend_yield = yahoo_data["dividend_yield"]
    shares_outstanding = yahoo_data["shares_outstanding"]

    return {
        "stock_price": stock_price,
        "volatility": round(annual_vol * 100, 2),
        "vol_start_label": vol_start_label,
        "vol_end_label": vol_end_label,
        "hist_monthly": hist_monthly,
        "hist_daily": hist_daily,
        "median_daily_volume": median_volume,
        "liquidity_shares": liquidity_shares,
        "volume_start_date": volume_start,
        "volume_end_date": volume_end,
        "dividend_yield": dividend_yield,
        "dividend_per_share": dividend_per_share,
        "report_date": report_date,
        "shares_outstanding": shares_outstanding,
    }


# ──────────────────────────────────────────────
# docx 操作
# ──────────────────────────────────────────────

def insert_paragraph_after(paragraph, text, font_name="ＭＳ Ｐ明朝", font_size=11):
    new_p = deepcopy(paragraph._element)
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    run_elem = paragraph._element.makeelement(qn('w:r'), {})
    rPr = run_elem.makeelement(qn('w:rPr'), {})
    rFonts = rPr.makeelement(qn('w:rFonts'), {
        qn('w:ascii'): font_name,
        qn('w:eastAsia'): font_name,
        qn('w:hAnsi'): font_name,
    })
    rPr.append(rFonts)
    sz = rPr.makeelement(qn('w:sz'), {qn('w:val'): str(font_size * 2)})
    szCs = rPr.makeelement(qn('w:szCs'), {qn('w:val'): str(font_size * 2)})
    rPr.append(sz)
    rPr.append(szCs)
    run_elem.append(rPr)
    t_elem = run_elem.makeelement(qn('w:t'), {})
    t_elem.text = text
    run_elem.append(t_elem)
    new_p.append(run_elem)
    paragraph._element.addnext(new_p)
    return new_p


def build_volatility_excel(hist_monthly, company_name):
    """yfinance月次データからボラティリティ計算Excelを生成"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = "ボラティリティ計算"

    title_font = Font(name="ＭＳ Ｐゴシック", size=14, bold=True)
    header_font = Font(name="ＭＳ Ｐゴシック", size=11, bold=True)
    data_font = Font(name="ＭＳ Ｐゴシック", size=11)
    formula_font = Font(name="ＭＳ Ｐゴシック", size=11, color="0000CC")
    result_font = Font(name="ＭＳ Ｐゴシック", size=12, bold=True, color="CC0000")
    header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    result_fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    ws.merge_cells("A1:E1")
    ws["A1"] = f"ボラティリティの計算過程（{company_name}）"
    ws["A1"].font = title_font

    headers = [("A", "対象月"), ("B", "株価"), ("C", "対数株価"), ("D", "対数収益率")]
    for col, label in headers:
        cell = ws[f"{col}3"]
        cell.value = label
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center")

    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 14

    # データ行（降順）
    data_rows = []
    for idx, row in hist_monthly.iterrows():
        date_str = idx.strftime("%Y/%m")
        price = round(float(row["Close"]), 2)
        data_rows.append((date_str, price))
    data_rows.sort(key=lambda x: x[0], reverse=True)

    n = len(data_rows)
    for i, (dt, price) in enumerate(data_rows):
        r = 4 + i
        ws[f"A{r}"] = dt
        ws[f"A{r}"].font = data_font
        ws[f"A{r}"].border = thin_border
        ws[f"B{r}"] = price
        ws[f"B{r}"].font = data_font
        ws[f"B{r}"].border = thin_border
        ws[f"B{r}"].number_format = "#,##0.00"
        ws[f"C{r}"] = f"=LN(B{r})"
        ws[f"C{r}"].font = formula_font
        ws[f"C{r}"].border = thin_border
        ws[f"C{r}"].number_format = "0.0000000"
        if i == n - 1:
            ws[f"D{r}"] = ""
        else:
            ws[f"D{r}"] = f"=C{r}-C{r+1}"
            ws[f"D{r}"].number_format = "0.000%"
        ws[f"D{r}"].font = formula_font if i < n - 1 else data_font
        ws[f"D{r}"].border = thin_border

    end_r = 3 + n
    ret_range = f"D4:D{end_r - 1}"
    calc_r = end_r + 2
    ws[f"A{calc_r}"] = "【ボラティリティ計算】"
    ws[f"A{calc_r}"].font = header_font

    r1 = calc_r + 1
    ws[f"A{r1}"] = "月次σ"
    ws[f"A{r1}"].font = data_font
    ws[f"B{r1}"] = f"=STDEVP({ret_range})"
    ws[f"B{r1}"].font = result_font
    ws[f"B{r1}"].number_format = "0.000000"
    ws[f"B{r1}"].fill = result_fill
    ws[f"B{r1}"].border = thin_border

    r2 = r1 + 1
    ws[f"A{r2}"] = "年率σ"
    ws[f"A{r2}"].font = data_font
    ws[f"B{r2}"] = f"=B{r1}*SQRT(12)"
    ws[f"B{r2}"].font = result_font
    ws[f"B{r2}"].number_format = "0.00%"
    ws[f"B{r2}"].fill = result_fill
    ws[f"B{r2}"].border = thin_border

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_volume_excel(hist_daily, company_name):
    """yfinance日次データから出来高中央値Excelを生成"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = "出来高データ"

    title_font = Font(name="ＭＳ Ｐゴシック", size=14, bold=True)
    header_font = Font(name="ＭＳ Ｐゴシック", size=11, bold=True)
    data_font = Font(name="ＭＳ Ｐゴシック", size=11)
    result_font = Font(name="ＭＳ Ｐゴシック", size=12, bold=True, color="CC0000")
    header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    result_fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    volumes = [int(row["Volume"]) for _, row in hist_daily.iterrows()]
    median_vol = int(np.median(volumes))
    liquidity = math.ceil(median_vol * 0.1)

    dates = [idx.strftime("%Y/%m/%d") for idx in hist_daily.index]
    start_d = dates[-1] if dates else ""
    end_d = dates[0] if dates else ""

    ws["A1"] = f"日次売買高の中央値（{company_name}）"
    ws["A1"].font = title_font
    ws["B1"] = median_vol
    ws["B1"].font = result_font
    ws["B1"].fill = result_fill
    ws["C1"] = "株"
    ws["C1"].font = result_font
    ws["A2"] = f"対象期間: {start_d}～{end_d}"
    ws["A2"].font = data_font

    ws["A4"] = "株価データ"
    ws["A4"].font = header_font

    headers = [("A", "企業・業界"), ("B", f"{company_name}")]
    ws["A6"] = "企業・業界"
    ws["A6"].font = header_font
    ws["B6"] = company_name
    ws["B6"].font = data_font

    col_headers = [("A", "日付"), ("B", "終値(調整後)"), ("C", "指数"), ("D", "出来高(調整後)")]
    for col, label in col_headers:
        c = ws[f"{col}9"]
        c.value = label
        c.font = header_font
        c.fill = header_fill
        c.border = thin_border

    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["D"].width = 16

    for i, (idx, row) in enumerate(hist_daily.iloc[::-1].iterrows()):
        r = 10 + i
        ws[f"A{r}"] = idx.strftime("%Y/%m/%d")
        ws[f"A{r}"].font = data_font
        ws[f"A{r}"].border = thin_border
        ws[f"B{r}"] = round(float(row["Close"]), 2)
        ws[f"B{r}"].font = data_font
        ws[f"B{r}"].border = thin_border
        ws[f"B{r}"].number_format = "#,##0.00"
        ws[f"D{r}"] = int(row["Volume"])
        ws[f"D{r}"].font = data_font
        ws[f"D{r}"].border = thin_border
        ws[f"D{r}"].number_format = "#,##0"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_bond_excel(all_bonds, eval_dt, exercise_end_dt, selected_name):
    """JSDA長期国債データからリスクフリーレートExcelを生成（JSDAフォーマット準拠）"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "売買参考統計値（長期国債）"

    title_font = Font(name="ＭＳ Ｐゴシック", size=12, bold=True)
    header_font = Font(name="ＭＳ Ｐゴシック", size=9, bold=True)
    data_font = Font(name="ＭＳ Ｐゴシック", size=9)
    selected_font = Font(name="ＭＳ Ｐゴシック", size=9, bold=True, color="CC0000")
    info_font = Font(name="ＭＳ Ｐゴシック", size=10)
    header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    selected_fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    right_align = Alignment(horizontal="right")

    # タイトル
    ws.merge_cells("A1:L1")
    ws["A1"] = "公社債店頭売買参考統計値（長期国債抜粋）"
    ws["A1"].font = title_font

    date_str = f"{eval_dt.year}/{eval_dt.month:02d}/{eval_dt.day:02d}"
    ws["A2"] = f"基準日: {date_str}　　権利行使期間終了日: {exercise_end_dt.year}/{exercise_end_dt.month:02d}/{exercise_end_dt.day:02d}"
    ws["A2"].font = info_font
    ws.merge_cells("A2:L2")

    # ヘッダー行1（カテゴリ）
    hr = 4
    ws.merge_cells(f"A{hr}:A{hr+1}")
    ws[f"A{hr}"] = "銘柄名"
    ws[f"A{hr}"].font = header_font
    ws[f"A{hr}"].fill = header_fill
    ws[f"A{hr}"].border = thin_border
    ws[f"A{hr}"].alignment = center

    ws.merge_cells(f"B{hr}:B{hr+1}")
    ws[f"B{hr}"] = "償還期日"
    ws[f"B{hr}"].font = header_font
    ws[f"B{hr}"].fill = header_fill
    ws[f"B{hr}"].border = thin_border
    ws[f"B{hr}"].alignment = center

    ws.merge_cells(f"C{hr}:C{hr+1}")
    ws[f"C{hr}"] = "利率"
    ws[f"C{hr}"].font = header_font
    ws[f"C{hr}"].fill = header_fill
    ws[f"C{hr}"].border = thin_border
    ws[f"C{hr}"].alignment = center

    # 平均値（D-G）
    ws.merge_cells(f"D{hr}:G{hr}")
    ws[f"D{hr}"] = "平均値"
    ws[f"D{hr}"].font = header_font
    ws[f"D{hr}"].fill = header_fill
    ws[f"D{hr}"].border = thin_border
    ws[f"D{hr}"].alignment = center

    # 中央値（H-K）
    ws.merge_cells(f"H{hr}:K{hr}")
    ws[f"H{hr}"] = "中央値"
    ws[f"H{hr}"].font = header_font
    ws[f"H{hr}"].fill = header_fill
    ws[f"H{hr}"].border = thin_border
    ws[f"H{hr}"].alignment = center

    # 満期差（L）
    ws.merge_cells(f"L{hr}:L{hr+1}")
    ws[f"L{hr}"] = "満期差\n(日)"
    ws[f"L{hr}"].font = header_font
    ws[f"L{hr}"].fill = header_fill
    ws[f"L{hr}"].border = thin_border
    ws[f"L{hr}"].alignment = center

    # ヘッダー行2（サブカラム）
    sub_headers = [
        ("D", "単価"), ("E", "前日比\n(銭)"), ("F", "複利利回り\n(%)"), ("G", "単利利回り\n(%)"),
        ("H", "単価"), ("I", "前日比\n(銭)"), ("J", "複利利回り\n(%)"), ("K", "単利利回り\n(%)"),
    ]
    for col, label in sub_headers:
        cell = ws[f"{col}{hr+1}"]
        cell.value = label
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center

    # 空セルにもボーダーを付ける
    for col in ["A", "B", "C", "L"]:
        ws[f"{col}{hr+1}"].border = thin_border
        ws[f"{col}{hr+1}"].fill = header_fill

    # カラム幅
    ws.column_dimensions["A"].width = 16
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 6
    for col in ["D", "H"]:
        ws.column_dimensions[col].width = 10
    for col in ["E", "I"]:
        ws.column_dimensions[col].width = 8
    for col in ["F", "G", "J", "K"]:
        ws.column_dimensions[col].width = 10
    ws.column_dimensions["L"].width = 8

    # データ行
    data_start = hr + 2
    for i, bond in enumerate(all_bonds):
        r = data_start + i
        is_selected = bond["name"] == selected_name
        font = selected_font if is_selected else data_font
        fill = selected_fill if is_selected else None

        def set_cell(col, val, fmt=None):
            cell = ws[f"{col}{r}"]
            cell.value = val
            cell.font = font
            cell.border = thin_border
            if fill:
                cell.fill = fill
            if fmt:
                cell.number_format = fmt

        set_cell("A", bond["name"])
        set_cell("B", bond["maturity"].strftime("%Y/%m/%d"))
        set_cell("C", bond.get("coupon", ""), "0.0")
        set_cell("D", bond.get("avg_price", ""), "0.00")
        set_cell("E", bond.get("avg_change", ""), "0")
        set_cell("F", bond.get("avg_compound", ""), "0.000")
        set_cell("G", bond.get("avg_simple", ""), "0.000")
        set_cell("H", bond.get("med_price", ""), "0.00")
        set_cell("I", bond.get("med_change", ""), "0")
        set_cell("J", bond.get("med_compound", ""), "0.000")
        set_cell("K", bond.get("med_simple", ""), "0.000")
        set_cell("L", bond["diff_days"], "#,##0")

    # 選定結果
    end_r = data_start + len(all_bonds) + 1
    ws.merge_cells(f"A{end_r}:L{end_r}")
    ws[f"A{end_r}"] = f"【選定】権利行使期間終了日に最も償還日が近い長期国債: {selected_name}（中央値 複利利回り {[b['yield_value'] for b in all_bonds if b['name'] == selected_name][0] if any(b['name'] == selected_name for b in all_bonds) else ''}%）"
    ws[f"A{end_r}"].font = Font(name="ＭＳ Ｐゴシック", size=10, bold=True, color="CC0000")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def replace_in_runs(paragraph, old_text, new_text):
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False
    new_full = full_text.replace(old_text, new_text)
    runs = paragraph.runs
    if not runs:
        return False
    remaining = new_full
    for run in runs:
        run_len = len(run.text)
        if len(remaining) <= 0:
            run.text = ""
        elif len(remaining) <= run_len:
            run.text = remaining
            remaining = ""
        else:
            run.text = remaining[:run_len]
            remaining = remaining[run_len:]
    if remaining:
        runs[-1].text += remaining
    return True


def replace_in_document(doc, old_text, new_text):
    count = 0
    for para in doc.paragraphs:
        if replace_in_runs(para, old_text, new_text):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_runs(para, old_text, new_text):
                        count += 1
    return count


def fmt_date_jp(dt):
    return f"{dt.year}年{dt.month}月{dt.day}日"


# ──────────────────────────────────────────────
# Handler
# ──────────────────────────────────────────────

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(content_length))
            ticker_code = body["ticker_code"]
            eval_date = body["eval_date"]
            eval_dt = datetime.strptime(eval_date, "%Y-%m-%d")

            # TOPページ入力項目
            exercise_start = body.get("exercise_start", "")
            exercise_end = body.get("exercise_end", "")
            assignee = body.get("assignee", "")
            resolution_date = body.get("resolution_date", "")
            warrant_total = body.get("warrant_total", "").replace(",", "")
            issuable_shares = body.get("issuable_shares", "").replace(",", "")
            fair_value_str = body.get("fair_value_per_share", "")
            special_terms = body.get("special_terms", "")
            market_risk_premium = body.get("market_risk_premium", "")
            default_rate = body.get("default_rate", "")
            credit_cost_input = body.get("credit_cost", "")
            bond_name = body.get("bond_name", "").strip()
            bond_maturity = body.get("bond_maturity", "").strip()
            bond_yield = body.get("bond_yield", "").strip()
            beta_value = body.get("beta", "")
            volatility_override = body.get("volatility", "")
            vol_start_override = body.get("vol_start_label", "")
            vol_end_override = body.get("vol_end_label", "")
            median_volume_override = body.get("median_volume", "")
            volume_start_override = body.get("volume_start", "")
            volume_end_override = body.get("volume_end", "")

            fair_value_per_share = float(fair_value_str) if fair_value_str else None

            # データ取得
            company_name_jp = fetch_japanese_company_name(ticker_code)
            profile = fetch_company_profile(ticker_code)
            data = fetch_stock_data(ticker_code, eval_date, exercise_end)

            # 国債データ自動取得
            jsda_all_bonds = []
            ex_end_dt = None
            if exercise_end:
                ex_end_dt = datetime.strptime(exercise_end, "%Y-%m-%d")
                try:
                    jsda = fetch_jsda_bond(eval_dt, ex_end_dt)
                    jsda_all_bonds = jsda.get("all_bonds", [])
                    if not bond_yield and jsda["yield_value"]:
                        bond_name = bond_name or jsda["name"]
                        bond_maturity = bond_maturity or jsda["maturity"]
                        bond_yield = jsda["yield_value"]
                except Exception:
                    pass

            # テンプレート読み込み
            if not os.path.exists(TEMPLATE_PATH):
                raise FileNotFoundError(f"テンプレートが見つかりません: {TEMPLATE_PATH}")
            doc = Document(TEMPLATE_PATH)

            # 置換
            # SPEEDAデータからの上書き
            vol_pct = float(volatility_override) if volatility_override else data['volatility']
            vol_start_lbl = vol_start_override if vol_start_override else data['vol_start_label']
            vol_end_lbl = vol_end_override if vol_end_override else data['vol_end_label']

            # 出来高: SPEEDAデータ優先
            if median_volume_override:
                median_vol = int(median_volume_override)
                liquidity = math.ceil(median_vol * 0.1)
            else:
                median_vol = data['median_daily_volume']
                liquidity = data['liquidity_shares']

            if volume_start_override and volume_end_override:
                vol_s_dt = datetime.strptime(volume_start_override, "%Y-%m-%d")
                vol_e_dt = datetime.strptime(volume_end_override, "%Y-%m-%d")
                volume_period = f"{fmt_date_jp(vol_s_dt)}から{fmt_date_jp(vol_e_dt)}"
            else:
                volume_period = f"{fmt_date_jp(data['volume_start_date'])}から{fmt_date_jp(data['volume_end_date'])}"

            # 権利行使期間（他の日付置換より先に実行）
            replacements = []
            if exercise_start and exercise_end:
                ex_start_dt = datetime.strptime(exercise_start, "%Y-%m-%d")
                ex_end_dt = datetime.strptime(exercise_end, "%Y-%m-%d")
                replacements.append(("2026年3月3日-", f"{fmt_date_jp(ex_start_dt)}-"))
                replacements.append(("2026年3月4日", fmt_date_jp(ex_end_dt)))

            replacements += [
                ("ジェリービーンズグループ", company_name_jp),
                ("3070", ticker_code),
                ("110円", f"{data['stock_price']:,}円"),
                ("62.54%", f"{vol_pct}%"),
                ("2021年2月- 2026年2月",
                 f"{vol_start_lbl}- {vol_end_lbl}"),
                ("2021年3月3日から2026年3月2日", volume_period),
                ("1,483,123", f"{median_vol:,}"),
                ("148,313", f"{liquidity:,}"),
                ("0%（0円/株）",
                 f"{data['dividend_yield']}%（{data['dividend_per_share']}円/株）"),
                ("2026年3月2日", fmt_date_jp(eval_dt)),
                ("79,440,000", f"{data['shares_outstanding']:,}"),
                ("宮崎明", profile['representative'].replace("\u3000", "")),
                ("宮崎\u3000明", profile['representative']),
                ("東京都台東区上野1-16-5", profile['address']),
                ("1990年4月", re.sub(r'\d{1,2}日$', '', profile['established'])),
                ("1月末", profile['settlement'].replace("日", "")),
            ]

            # 売買参考統計値から抽出された国債情報
            if bond_name:
                replacements.append(("長期国債362", bond_name))
            if bond_maturity:
                bm_dt = datetime.strptime(bond_maturity, "%Y-%m-%d")
                replacements.append(("2031年3月20日", fmt_date_jp(bm_dt)))
            if bond_yield:
                replacements.append(("1.591%", f"{bond_yield}%"))

            # CAPM各変数（テーブル内の個別セル）
            if market_risk_premium:
                replacements.append(("9.3%", f"{market_risk_premium}%"))

            # 対指数β
            if beta_value:
                replacements.append(("0.567", str(beta_value)))

            # デフォルト率
            default_rate_num = float(default_rate) if default_rate else 17.92
            if default_rate and default_rate != "17.92":
                replacements.append(("17.92%", f"{default_rate_num}%"))
                # 回収率: max(59.1 - 8.356 * デフォルト率, 0)
                recovery = max(59.1 - 8.356 * default_rate_num, 0)
                recovery_str = f"{recovery:.1f}%" if recovery > 0 else "0%"
                replacements.append(("0%\n", f"{recovery_str}\n"))

            # クレジットコスト
            if credit_cost_input and credit_cost_input != "21.83":
                replacements.append(("21.83%", f"{credit_cost_input}%"))

            # CAPM計算式の自動計算
            rfr = float(bond_yield) if bond_yield else 1.591
            mrp = float(market_risk_premium) if market_risk_premium else 9.3
            beta_num = float(beta_value) if beta_value else 0.567
            credit_cost = float(credit_cost_input) if credit_cost_input else 21.83
            capm_result = round(rfr + mrp * beta_num + credit_cost, 2)
            replacements.append((
                "= 1.591% + 9.3%\u00d7 0.567 + 21.83%",
                f"= {rfr}% + {mrp}%\u00d7 {beta_num} + {credit_cost}%"
            ))
            replacements.append(("= 28.69%", f"= {capm_result}%"))

            # ●プレースホルダー（テーブル1）
            if resolution_date:
                try:
                    res_dt = datetime.strptime(resolution_date, "%Y-%m-%d")
                    replacements.append(("2026年●月●日", fmt_date_jp(res_dt)))
                except ValueError:
                    # 「未定」等のテキストの場合はそのまま置換
                    replacements.append(("2026年●月●日", resolution_date))
            if warrant_total:
                replacements.append(("●個", f"{int(warrant_total):,}個"))
            if issuable_shares:
                replacements.append(("●株", f"{int(issuable_shares):,}株"))
            # 行使による払込価額 = 株価と同額
            replacements.append(("●円", f"{data['stock_price']:,}円"))

            # 公正価値 → 株価比率
            if fair_value_per_share is not None:
                price_ratio = round(fair_value_per_share / data['stock_price'] * 100, 2)
                fair_value_per_unit = round(fair_value_per_share * 100)
                replacements.append(("公正価値113円", f"公正価値{fair_value_per_unit:,}円"))
                replacements.append(("1.13円/株", f"{fair_value_per_share}円/株"))
                replacements.append(("当初株価の1.03%", f"当初株価の{price_ratio:.2f}%"))

            for old, new in replacements:
                replace_in_document(doc, old, new)

            # 割当先（Table1 R1 C1）- 改行対応
            if assignee:
                try:
                    cell = doc.tables[1].rows[1].cells[1]
                    # 既存のrunをクリア
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ""
                    lines = assignee.split("\n")
                    first_para = cell.paragraphs[0]
                    if first_para.runs:
                        first_para.runs[0].text = lines[0]
                        first_para.runs[0].font.size = Pt(11)
                        first_para.runs[0].font.name = "ＭＳ Ｐ明朝"
                        rPr = first_para.runs[0]._element.find(qn('w:rPr'))
                        if rPr is not None:
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = rPr.makeelement(qn('w:rFonts'), {})
                                rPr.insert(0, rFonts)
                            rFonts.set(qn('w:eastAsia'), "ＭＳ Ｐ明朝")
                    else:
                        run = first_para.add_run(lines[0])
                        run.font.size = Pt(11)
                        run.font.name = "ＭＳ Ｐ明朝"
                    for line in lines[1:]:
                        insert_paragraph_after(cell.paragraphs[-1], line)
                except Exception:
                    pass

            # 査定に関連する特約条項（Table1 R6 C1）
            if special_terms:
                try:
                    cell = doc.tables[1].rows[6].cells[1]
                    # 既存段落・runを全てクリア
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ""
                    # 最初の段落に1行目を設定（フォント11pt保証）
                    lines = special_terms.split("\n")
                    first_para = cell.paragraphs[0]
                    if first_para.runs:
                        first_para.runs[0].text = lines[0]
                        first_para.runs[0].font.size = Pt(11)
                        first_para.runs[0].font.name = "ＭＳ Ｐ明朝"
                        rPr = first_para.runs[0]._element.find(qn('w:rPr'))
                        if rPr is not None:
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = rPr.makeelement(qn('w:rFonts'), {})
                                rPr.insert(0, rFonts)
                            rFonts.set(qn('w:eastAsia'), "ＭＳ Ｐ明朝")
                    else:
                        # runがない場合は新規作成
                        run = first_para.add_run(lines[0])
                        run.font.size = Pt(11)
                        run.font.name = "ＭＳ Ｐ明朝"
                        rPr = run._element.find(qn('w:rPr'))
                        if rPr is not None:
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is not None:
                                rFonts.set(qn('w:eastAsia'), "ＭＳ Ｐ明朝")
                    # 残りの行は新規段落追加（font_size=11がデフォルト）
                    for line in lines[1:]:
                        insert_paragraph_after(cell.paragraphs[-1], line)
                except Exception:
                    pass

            # 報告書を保存
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name
            with open(tmp_path, "rb") as f:
                docx_bytes = f.read()
            os.unlink(tmp_path)

            eval_ym = eval_dt.strftime("%Y%m")
            docx_filename = f"{eval_ym}_新株予約権評価報告書_株式会社{company_name_jp}.docx"

            # ボラティリティ・出来高・国債のExcelを生成
            vol_excel = build_volatility_excel(data["hist_monthly"], company_name_jp)
            volume_excel = build_volume_excel(data["hist_daily"], company_name_jp)

            # ZIPにまとめて返す
            import zipfile
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                zf.writestr(docx_filename, docx_bytes)
                zf.writestr(f"{company_name_jp}_ボラティリティ計算.xlsx", vol_excel)
                zf.writestr(f"{company_name_jp}_出来高中央値.xlsx", volume_excel)
                if jsda_all_bonds and ex_end_dt:
                    bond_excel = build_bond_excel(
                        jsda_all_bonds, eval_dt, ex_end_dt, bond_name)
                    zf.writestr(f"{company_name_jp}_リスクフリーレート.xlsx", bond_excel)
            zip_bytes = zip_buf.getvalue()

            zip_filename = f"{eval_ym}_株式会社{company_name_jp}_算定資料.zip"

            self.send_response(200)
            self.send_header("Content-Type", "application/zip")
            self.send_header("Content-Disposition",
                             f"attachment; filename*=UTF-8''{urllib.parse.quote(zip_filename)}")
            self.send_header("Content-Length", str(len(zip_bytes)))
            self.end_headers()
            self.wfile.write(zip_bytes)

        except Exception as e:
            self.send_response(500)
            self.send_header("Content-Type", "application/json")
            self.end_headers()
            self.wfile.write(json.dumps({"error": str(e)}).encode())

#!/usr/bin/env python3
"""
新株予約権価値算定報告書 自動生成スクリプト

使い方:
  python3 generate_report.py <銘柄コード> <評価基準日>
  python3 generate_report.py 3070 2025-06-13

自動取得するデータ:
  - 会社名 (日本語)
  - 株価 (S) : 評価基準日の終値
  - ボラティリティ (σ) : 評価基準日から過去5年の月次株価
  - 日次売買高中央値 → 流動性 (10%)
  - 配当率 (q)
  - 各種日付 (報告書日付、算定基準日、ボラティリティ算出期間、売買高期間)
"""

import sys
import os
import re
import math
import urllib.request
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
from docx import Document
import yfinance as yf
import numpy as np
import pandas as pd
import warnings
warnings.filterwarnings("ignore")

TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "template.docx")
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")


# ──────────────────────────────────────────────
# 1. データ取得
# ──────────────────────────────────────────────

def fetch_japanese_company_name(ticker_code: str) -> str:
    """Yahoo Finance Japan から日本語社名を取得 (「株式会社」「(株)」なし)"""
    url = f"https://finance.yahoo.co.jp/quote/{ticker_code}.T"
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(req, timeout=10) as resp:
        html = resp.read().decode("utf-8")
    m = re.search(r"<title>(.*?)【\d+】", html)
    if not m:
        raise ValueError(f"社名を取得できませんでした: {ticker_code}")
    name = m.group(1).strip()
    # 「(株)」「株式会社」を除去
    name = re.sub(r"^\(株\)", "", name)
    name = re.sub(r"\(株\)$", "", name)
    name = re.sub(r"^株式会社", "", name)
    name = re.sub(r"株式会社$", "", name)
    return name.strip()


def fetch_company_profile(ticker_code: str) -> dict:
    """Yahoo Finance Japan から会社プロフィール情報を取得"""
    url = f"https://finance.yahoo.co.jp/quote/{ticker_code}.T/profile"
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(req, timeout=10) as resp:
        html = resp.read().decode("utf-8")

    def extract(label):
        m = re.search(rf'<th[^>]*>{label}</th>\s*<td[^>]*>(.*?)</td>', html, re.DOTALL)
        if m:
            return re.sub(r'<[^>]+>', '', m.group(1)).strip()
        return ""

    # 所在地（郵便番号除去）
    address = ""
    m = re.search(r'〒[\d\-]+\s*(.+?)(?=<|\")', html)
    if m:
        address = m.group(1).strip()

    # 代表者名から肩書と名前を分離
    representative = extract("代表者名")  # 例: "宮崎　明"

    return {
        "representative": representative,
        "address": address,
        "established": extract("設立年月日"),  # 例: "1990年4月10日"
        "settlement": extract("決算"),  # 例: "1月末日"
    }


def fetch_stock_data(ticker_code: str, eval_date: str):
    """
    yfinance から株価データを取得

    Args:
        ticker_code: 銘柄コード (例: "3070")
        eval_date: 評価基準日 "YYYY-MM-DD" (例: "2025-06-13")

    Returns:
        dict
    """
    ticker = yf.Ticker(f"{ticker_code}.T")
    eval_dt = datetime.strptime(eval_date, "%Y-%m-%d")

    # --- 株価 (評価基準日の終値) ---
    # 評価基準日前後の数日分を取得
    start = (eval_dt - timedelta(days=10)).strftime("%Y-%m-%d")
    end = (eval_dt + timedelta(days=1)).strftime("%Y-%m-%d")
    hist_around = ticker.history(start=start, end=end)
    # 評価基準日以前の最新終値
    hist_before = hist_around[hist_around.index.strftime("%Y-%m-%d") <= eval_date]
    if len(hist_before) == 0:
        raise ValueError(f"評価基準日 {eval_date} の株価データが取得できません")
    stock_price = int(hist_before["Close"].iloc[-1])
    actual_price_date = hist_before.index[-1]

    # --- ボラティリティ (評価基準日前月から過去5年の月次) ---
    # テンプレート: "2020年5月- 2025年5月の月次株価"
    # → 評価基準日の前月末から5年前の同月の月次データ
    vol_end_month = eval_dt.replace(day=1) - timedelta(days=1)  # 前月末
    vol_start_month = vol_end_month - relativedelta(years=5)
    vol_start = vol_start_month.replace(day=1).strftime("%Y-%m-%d")
    vol_end = (vol_end_month + timedelta(days=1)).strftime("%Y-%m-%d")

    hist_monthly = ticker.history(start=vol_start, end=vol_end, interval="1mo")
    returns = np.log(hist_monthly["Close"] / hist_monthly["Close"].shift(1)).dropna()
    monthly_vol = returns.std()
    annual_vol = monthly_vol * np.sqrt(12)

    vol_start_label = vol_start_month.strftime("%Y年%-m月")
    vol_end_label = vol_end_month.strftime("%Y年%-m月")

    # --- 日次売買高の中央値 (評価基準日前日から過去5年) ---
    # テンプレート: "2020年6月13日から2025年6月12日"
    report_date = eval_dt - timedelta(days=1)  # 報告書日 = 基準日前日
    volume_end = report_date
    volume_start = volume_end - relativedelta(years=5)

    hist_daily = ticker.history(
        start=volume_start.strftime("%Y-%m-%d"),
        end=(volume_end + timedelta(days=1)).strftime("%Y-%m-%d"),
    )
    median_volume = int(hist_daily["Volume"].median())
    liquidity_shares = math.ceil(median_volume * 0.1)  # 10% 切り上げ

    # --- 配当 ---
    dividends = ticker.dividends
    if len(dividends) > 0:
        # 評価基準日以前の直近配当
        divs_before = dividends[dividends.index.strftime("%Y-%m-%d") <= eval_date]
        if len(divs_before) > 0:
            recent_div = divs_before.iloc[-1]
            dividend_per_share = int(recent_div)
        else:
            dividend_per_share = 0
    else:
        dividend_per_share = 0

    dividend_yield = round((dividend_per_share / stock_price * 100), 2) if stock_price > 0 else 0.0

    # --- 発行済株式総数 ---
    shares_outstanding = ticker.info.get("sharesOutstanding", 0)

    return {
        "stock_price": stock_price,
        "stock_price_date": actual_price_date,
        "volatility": round(annual_vol * 100, 2),
        "vol_start_label": vol_start_label,
        "vol_end_label": vol_end_label,
        "median_daily_volume": median_volume,
        "liquidity_shares": liquidity_shares,
        "volume_start_date": volume_start,
        "volume_end_date": volume_end,
        "dividend_yield": dividend_yield,
        "dividend_per_share": dividend_per_share,
        "report_date": report_date,
        "shares_outstanding": shares_outstanding,
    }


# ──────────────────────────────────────────────
# 2. docx テキスト置換 (フォーマット保持)
# ──────────────────────────────────────────────

def insert_paragraph_after(paragraph, text, font_name="ＭＳ Ｐ明朝"):
    """指定段落の直後に新しい段落を挿入 (同じフォーマットで)"""
    from docx.oxml.ns import qn
    from copy import deepcopy
    new_p = deepcopy(paragraph._element)
    # テキストをクリアして新しいテキストを設定
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    run_elem = paragraph._element.makeelement(qn('w:r'), {})
    rPr = run_elem.makeelement(qn('w:rPr'), {})
    rFonts = rPr.makeelement(qn('w:rFonts'), {
        qn('w:ascii'): font_name,
        qn('w:eastAsia'): font_name,
        qn('w:hAnsi'): font_name,
    })
    rPr.append(rFonts)
    run_elem.append(rPr)
    t_elem = run_elem.makeelement(qn('w:t'), {})
    t_elem.text = text
    run_elem.append(t_elem)
    new_p.append(run_elem)
    paragraph._element.addnext(new_p)
    return new_p


def replace_in_runs(paragraph, old_text, new_text):
    """
    段落内の複数 run にまたがるテキストを置換。
    フォーマットは各 run のものを保持。
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False

    new_full = full_text.replace(old_text, new_text)
    runs = paragraph.runs
    if len(runs) == 0:
        return False

    # テキストを再配分
    remaining = new_full
    for run in runs:
        run_len = len(run.text)
        if len(remaining) <= 0:
            run.text = ""
        elif len(remaining) <= run_len:
            run.text = remaining
            remaining = ""
        else:
            run.text = remaining[:run_len]
            remaining = remaining[run_len:]

    if remaining:
        runs[-1].text += remaining

    return True


def replace_in_document(doc, old_text, new_text):
    """ドキュメント全体 (段落 + テーブル) でテキスト置換"""
    count = 0
    for para in doc.paragraphs:
        if replace_in_runs(para, old_text, new_text):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_runs(para, old_text, new_text):
                        count += 1
    return count


# ──────────────────────────────────────────────
# 3. 日付フォーマット
# ──────────────────────────────────────────────

def fmt_date_jp(dt):
    """datetime → "2025年6月13日" """
    return f"{dt.year}年{dt.month}月{dt.day}日"


def fmt_year_month_jp(dt):
    """datetime → "2025年6月" """
    return f"{dt.year}年{dt.month}月"


# ──────────────────────────────────────────────
# 4. メイン処理
# ──────────────────────────────────────────────

def main():
    if len(sys.argv) < 3:
        print("使い方: python3 generate_report.py <銘柄コード> <評価基準日> [1株あたり公正価値]")
        print("例:     python3 generate_report.py 3070 2025-06-13 51.04")
        sys.exit(1)

    ticker_code = sys.argv[1]
    eval_date = sys.argv[2]
    eval_dt = datetime.strptime(eval_date, "%Y-%m-%d")
    fair_value_per_share = float(sys.argv[3]) if len(sys.argv) >= 4 else None

    print(f"銘柄コード: {ticker_code}")
    print(f"評価基準日: {eval_date}")
    print()

    # ── データ取得 ──
    print("日本語社名取得中...")
    company_name_jp = fetch_japanese_company_name(ticker_code)
    print(f"  社名: 株式会社{company_name_jp}")

    print("会社プロフィール取得中...")
    profile = fetch_company_profile(ticker_code)
    print(f"  代表者: {profile['representative']}")
    print(f"  所在地: {profile['address']}")
    print(f"  設立: {profile['established']}")
    print(f"  決算: {profile['settlement']}")

    print("株価データ取得中...")
    data = fetch_stock_data(ticker_code, eval_date)

    print(f"  株価 (S): {data['stock_price']}円")
    print(f"  ボラティリティ (σ): {data['volatility']}%")
    print(f"    算出期間: {data['vol_start_label']} - {data['vol_end_label']}")
    print(f"  日次売買高中央値: {data['median_daily_volume']:,}株")
    print(f"  流動性 (10%): {data['liquidity_shares']:,}株")
    print(f"    算出期間: {fmt_date_jp(data['volume_start_date'])} - {fmt_date_jp(data['volume_end_date'])}")
    dividend_per_share = data['dividend_per_share']
    dividend_yield = data['dividend_yield']
    print(f"  配当(Yahoo Finance): {dividend_per_share}円/株 ({dividend_yield}%)")
    print(f"  発行済株式総数: {data['shares_outstanding']:,}株")
    print(f"  報告書日付: {fmt_date_jp(data['report_date'])}")

    # ── テンプレート読み込み ──
    print("\nテンプレート読み込み中...")
    doc = Document(TEMPLATE_PATH)

    # ── 置換定義 ──
    # テンプレート内のジェリービーンズグループ固有の値 → 新しい値
    replacements = [
        # 社名 (「株式会社」は付いていない部分を置換)
        ("ジェリービーンズグループ", company_name_jp),

        # 銘柄コード
        ("3070", ticker_code),

        # 株価
        ("220円", f"{data['stock_price']}円"),

        # ボラティリティ
        ("34.47%", f"{data['volatility']}%"),

        # ボラティリティ算出期間 (テーブル注記)
        ("2020年5月- 2025年5月", f"{data['vol_start_label']}- {data['vol_end_label']}"),

        # 日次売買高中央値 (長い文字列を先に)
        ("24,600", f"{data['median_daily_volume']:,}"),

        # 流動性 10%
        ("2,460", f"{data['liquidity_shares']:,}"),

        # 売買高算出期間 (本文中 + SPEEDA注記、日付を含むので先に)
        ("2020年6月13日から2025年6月12日",
         f"{fmt_date_jp(data['volume_start_date'])}から{fmt_date_jp(data['volume_end_date'])}"),

        # 配当 (長い文字列を先に置換 → 短い文字列が壊れない)
        ("0%（0円/株）",
         f"{dividend_yield}%（{dividend_per_share}円/株）"),

        # 報告書日付 (表紙) → 評価基準日と同日
        # ※売買高期間の日付が先に置換されるため、残っている箇所のみ対象
        ("2025年6月12日", fmt_date_jp(eval_dt)),

        # 評価基準日 (本文)
        ("2025年6月13日", fmt_date_jp(eval_dt)),

        # 発行済株式総数
        ("33,950,000", f"{data['shares_outstanding']:,}"),

        # 代表者名 (本文中 "宮崎明" を置換、全角スペース含むパターンも)
        ("宮崎明", profile['representative'].replace("　", "")),
        ("宮崎\u3000明", profile['representative']),

        # 所在地
        ("東京都台東区上野1-16-5", profile['address']),

        # 設立年月 (テンプレート: "1990年4月")
        ("1990年4月", profile['established'].replace("10日", "").rstrip("日")),

        # 決算日 (テンプレート: "1月末")
        ("1月末", profile['settlement'].replace("日", "")),

        # 権利行使価格 = 株価と同額
        ("95円", f"{data['stock_price']}円"),
    ]

    # ── 1株あたり公正価値 → 株価比率計算 ──
    if fair_value_per_share is not None:
        stock_price = data['stock_price']
        # 株価比率 = 公正価値 / 株価 × 100 (小数第2位で切り上げ)
        import math
        price_ratio_raw = fair_value_per_share / stock_price * 100
        price_ratio = math.ceil(price_ratio_raw * 100) / 100  # 小数第2位切り上げ
        print(f"\n  1株あたり公正価値: {fair_value_per_share}円")
        print(f"  株価比率: {fair_value_per_share} / {stock_price} × 100 = {price_ratio_raw:.4f}% → 切上げ {price_ratio:.2f}%")

        # テンプレート: 「51.04円/株　当初株価の23.20%」を置換
        replacements.append(("51.04円/株", f"{fair_value_per_share}円/株"))
        replacements.append(("23.20%", f"{price_ratio:.2f}%"))

    print("\n置換実行中...")
    for old, new in replacements:
        n = replace_in_document(doc, old, new)
        if old == new:
            continue  # 同じ値の場合はスキップ表示
        status = "OK" if n > 0 else "--"
        print(f"  [{status}] 「{old}」→「{new}」({n}箇所)")

    # ── 保存 ──
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    eval_ym = eval_dt.strftime("%Y%m")
    output_filename = f"{eval_ym}_新株予約権評価報告書_株式会社{company_name_jp}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(output_path)
    print(f"\n生成完了: {output_path}")

    # ── 未対応変数の一覧 ──
    print("\n" + "=" * 50)
    print("【自動取得済み】")
    print(f"  社名: 株式会社{company_name_jp}")
    print(f"  株価 (S): {data['stock_price']}円")
    print(f"  ボラティリティ (σ): {data['volatility']}%")
    print(f"  売買高中央値: {data['median_daily_volume']:,}株 / 流動性: {data['liquidity_shares']:,}株")
    print(f"  配当: {data['dividend_per_share']}円/株")
    print(f"  報告書日付: {fmt_date_jp(data['report_date'])}")
    print(f"  評価基準日: {fmt_date_jp(eval_dt)}")

    print("\n【テンプレートのまま（今後対応）】")
    print("  - 権利行使価格 (K): 95円")
    print("  - 権利行使期間 (t): 2025年8月16日-2030年8月15日")
    print("  - リスクフリーレート (r): 1.027%")
    print("  - 代表者名: 宮崎明")
    print("  - 代表者肩書: 代表取締役社長")
    print("  - 所在地: 東京都台東区上野1-16-5")
    print("  - 事業内容")
    print("  - 設立年月")
    print("  - 決算日")
    print("  - 発行済株式総数: 33,950,000株")
    print("  - 新株予約権の総数: 618,750個")
    print("  - 決議年月日")
    print("  - 市場リスクプレミアム: 9.0%")
    print("  - β: 0.500")
    print("  - クレジットコスト: 22.86%")
    print("  - 公正価値: 5,104円")
    print("  - 希薄化率")


if __name__ == "__main__":
    main()
